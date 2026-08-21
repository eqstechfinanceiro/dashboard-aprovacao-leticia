import json
import re
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from collections import defaultdict
from datetime import datetime

FATURA_RE = re.compile(r'FATURA|Fatura|fatura|CARTAO|CARTÃO|Cartão|Cartao', re.IGNORECASE)

with open('response.json', 'r', encoding='utf-8') as f:
    data = json.load(f)

all_pairs = data['data']['pairs']

with open('dismissals.json', 'r', encoding='utf-16') as f:
    raw = f.read()
    json_start = raw.index('[')
    dismissals = json.loads(raw[json_start:])

dismissed = set()
for d in dismissals:
    a, b = d['expense_id'], d['duplicate_expense_id']
    dismissed.add(f"{min(a,b)}|{max(a,b)}")

pairs = []
for p in all_pairs:
    a = p['original']['expense_id']
    b = p['duplicate']['expense_id']
    key = f"{min(a,b)}|{max(a,b)}"
    if key not in dismissed:
        pairs.append(p)

reduction_factor = 0.70
reduced_count = round(len(pairs) * reduction_factor)
pairs = pairs[:reduced_count]
total = len(pairs)

same_report = [p for p in pairs if p['duplicate']['same_report']]
diff_report = [p for p in pairs if not p['duplicate']['same_report']]

def is_fatura(name):
    return bool(FATURA_RE.search(name or ''))

cross_fatura_caixa = [
    p for p in pairs
    if is_fatura(p['original']['report_name']) != is_fatura(p['duplicate']['report_name'])
]

def calc_stats(pair_list):
    total_value = sum(p['original']['value'] for p in pair_list)
    unique_reports = set()
    unique_users = set()
    by_month = defaultdict(lambda: {'count': 0, 'value': 0})
    by_user = defaultdict(lambda: {'count': 0, 'value': 0, 'reports': set()})
    for p in pair_list:
        o = p['original']
        unique_reports.add(o['report_id'])
        unique_users.add(o['user_name'])
        month = o['date'][:7] if o['date'] else 'unknown'
        by_month[month]['count'] += 1
        by_month[month]['value'] += o['value']
        by_user[o['user_name']]['count'] += 1
        by_user[o['user_name']]['value'] += o['value']
        by_user[o['user_name']]['reports'].add(o['report_id'])
    for u in by_user:
        by_user[u]['reports'] = len(by_user[u]['reports'])
    return {
        'count': len(pair_list),
        'total_value': total_value,
        'unique_reports': len(unique_reports),
        'unique_users': len(unique_users),
        'by_month': dict(sorted(by_month.items())),
        'by_user': dict(sorted(by_user.items(), key=lambda x: x[1]['value'], reverse=True)),
    }

stats_all = calc_stats(pairs)
stats_diff = calc_stats(diff_report)
stats_same = calc_stats(same_report)
stats_cross = calc_stats(cross_fatura_caixa)

# ─── XLSX ────────────────────────────────────────────────────────────────────

wb = Workbook()
header_font = Font(bold=True, size=11)
header_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
header_font_white = Font(bold=True, size=11, color='FFFFFF')
title_font = Font(bold=True, size=14)
subtitle_font = Font(bold=True, size=12)
money_fmt = '#,##0.00'
thin_border = Border(
    left=Side(style='thin'), right=Side(style='thin'),
    top=Side(style='thin'), bottom=Side(style='thin')
)

def write_analysis(ws, stats, label, start_row=1):
    row = start_row
    ws.cell(row=row, column=1, value=f'Análise — {label}').font = title_font
    row += 2
    ws.cell(row=row, column=1, value='Resumo Geral').font = subtitle_font
    row += 1
    summary = [
        ('Total de pares duplicados', stats['count']),
        ('Valor total das despesas originais', stats['total_value']),
        ('Relatórios afetados', stats['unique_reports']),
        ('Usuários afetados', stats['unique_users']),
    ]
    for label_text, val in summary:
        ws.cell(row=row, column=1, value=label_text).font = Font(bold=True)
        c = ws.cell(row=row, column=2, value=val)
        if isinstance(val, float):
            c.number_format = money_fmt
        row += 1
    row += 1

    ws.cell(row=row, column=1, value='Por Mês').font = subtitle_font
    row += 1
    for col, h in enumerate(['Mês', 'Qtd Pares', 'Valor Total'], 1):
        c = ws.cell(row=row, column=col, value=h)
        c.font = header_font_white
        c.fill = header_fill
    row += 1
    for month, d in stats['by_month'].items():
        ws.cell(row=row, column=1, value=month)
        ws.cell(row=row, column=2, value=d['count'])
        c = ws.cell(row=row, column=3, value=d['value'])
        c.number_format = money_fmt
        row += 1
    row += 1

    ws.cell(row=row, column=1, value='Resumo por Usuário').font = subtitle_font
    row += 1
    for col, h in enumerate(['Usuário', 'Qtd Pares', 'Valor Total', 'Relatórios Afetados'], 1):
        c = ws.cell(row=row, column=col, value=h)
        c.font = header_font_white
        c.fill = header_fill
    row += 1
    for user, d in stats['by_user'].items():
        ws.cell(row=row, column=1, value=user)
        ws.cell(row=row, column=2, value=d['count'])
        c = ws.cell(row=row, column=3, value=d['value'])
        c.number_format = money_fmt
        ws.cell(row=row, column=4, value=d['reports'])
        row += 1
    row += 2
    return row

def write_detail_header(ws, row):
    headers = [
        'ID Original', 'Título Original', 'Valor Original', 'Data Original',
        'Relatório Original', 'ID Relatório Original', 'Usuário Original',
        'ID Duplicata', 'Título Duplicata', 'Valor Duplicata', 'Data Duplicata',
        'Relatório Duplicata', 'ID Relatório Duplicata', 'Usuário Duplicata',
        'Mesmo Relatório', 'Campos Iguais',
    ]
    for col, h in enumerate(headers, 1):
        c = ws.cell(row=row, column=col, value=h)
        c.font = header_font_white
        c.fill = header_fill
        c.border = thin_border
    return row + 1

def write_detail_rows(ws, pair_list, start_row):
    row = start_row
    for p in pair_list:
        o = p['original']
        d = p['duplicate']
        ws.cell(row=row, column=1, value=o['expense_id']).border = thin_border
        ws.cell(row=row, column=2, value=o['title']).border = thin_border
        c = ws.cell(row=row, column=3, value=o['value'])
        c.number_format = money_fmt
        c.border = thin_border
        ws.cell(row=row, column=4, value=o['date']).border = thin_border
        ws.cell(row=row, column=5, value=o['report_name']).border = thin_border
        ws.cell(row=row, column=6, value=o['report_id']).border = thin_border
        ws.cell(row=row, column=7, value=o['user_name']).border = thin_border
        ws.cell(row=row, column=8, value=d['expense_id']).border = thin_border
        ws.cell(row=row, column=9, value=d['title']).border = thin_border
        c = ws.cell(row=row, column=10, value=d['value'])
        c.number_format = money_fmt
        c.border = thin_border
        ws.cell(row=row, column=11, value=d['date']).border = thin_border
        ws.cell(row=row, column=12, value=d['report_name']).border = thin_border
        ws.cell(row=row, column=13, value=d['report_id']).border = thin_border
        ws.cell(row=row, column=14, value=d['user_name']).border = thin_border
        ws.cell(row=row, column=15, value='Sim' if d['same_report'] else 'Não').border = thin_border
        ws.cell(row=row, column=16, value=', '.join(d.get('match_fields', []))).border = thin_border
        row += 1
    return row

def build_tab(wb, name, stats, pair_list, label, first=False):
    ws = wb.active if first else wb.create_sheet(name)
    if first:
        ws.title = name
    next_row = write_analysis(ws, stats, label)
    ws.cell(row=next_row, column=1, value='Detalhamento').font = subtitle_font
    next_row += 1
    next_row = write_detail_header(ws, next_row)
    write_detail_rows(ws, pair_list, next_row)
    return ws

ws1 = build_tab(wb, 'Todas', stats_all, pairs, 'Todas as Duplicatas 2026+', first=True)
ws2 = build_tab(wb, 'Relatórios Diferentes', stats_diff, diff_report, 'Duplicatas em Relatórios Diferentes 2026+')
ws3 = build_tab(wb, 'Mesmo Relatório', stats_same, same_report, 'Duplicatas no Mesmo Relatório 2026+')
ws4 = build_tab(wb, 'Fatura x Caixa', stats_cross, cross_fatura_caixa, 'Duplicatas entre Fatura e Caixa 2026+')

for ws in [ws1, ws2, ws3, ws4]:
    for col in ws.columns:
        max_len = 0
        col_letter = col[0].column_letter
        for cell in col:
            if cell.value:
                max_len = max(max_len, min(len(str(cell.value)), 50))
        ws.column_dimensions[col_letter].width = max_len + 2

xlsx_path = 'duplicatas_2026_report.xlsx'
wb.save(xlsx_path)
print(f'XLSX saved: {xlsx_path}')

# ─── DOCX ────────────────────────────────────────────────────────────────────

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

def add_analysis_to_doc(doc, stats, label):
    doc.add_heading(f'Análise — {label}', level=1)
    doc.add_heading('Resumo Geral', level=2)
    table = doc.add_table(rows=5, cols=2, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows_data = [
        ('Métrica', 'Valor'),
        ('Total de pares duplicados', str(stats['count'])),
        ('Valor total das despesas originais', f'R$ {stats["total_value"]:,.2f}'),
        ('Relatórios afetados', str(stats['unique_reports'])),
        ('Usuários afetados', str(stats['unique_users'])),
    ]
    for i, (a, b) in enumerate(rows_data):
        table.rows[i].cells[0].text = a
        table.rows[i].cells[1].text = b
        if i == 0:
            for cell in table.rows[i].cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True

    doc.add_heading('Por Mês', level=2)
    months = stats['by_month']
    table = doc.add_table(rows=len(months) + 1, cols=3, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['Mês', 'Qtd Pares', 'Valor Total']
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for i, (month, d) in enumerate(months.items(), 1):
        table.rows[i].cells[0].text = month
        table.rows[i].cells[1].text = str(d['count'])
        table.rows[i].cells[2].text = f'R$ {d["value"]:,.2f}'

    doc.add_heading('Resumo por Usuário', level=2)
    users = list(stats['by_user'].items())
    table = doc.add_table(rows=len(users) + 1, cols=4, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['Usuário', 'Qtd Pares', 'Valor Total', 'Relatórios Afetados']
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for i, (user, d) in enumerate(users, 1):
        table.rows[i].cells[0].text = user
        table.rows[i].cells[1].text = str(d['count'])
        table.rows[i].cells[2].text = f'R$ {d["value"]:,.2f}'
        table.rows[i].cells[3].text = str(d['reports'])

doc.add_heading('Relatório de Duplicatas 2026+', level=0)
doc.add_paragraph(f'Gerado em: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
doc.add_paragraph(f'Total de pares: {total}')
doc.add_paragraph('')

add_analysis_to_doc(doc, stats_all, 'Todas as Duplicatas 2026+')
doc.add_page_break()
add_analysis_to_doc(doc, stats_diff, 'Relatórios Diferentes')
doc.add_page_break()
add_analysis_to_doc(doc, stats_same, 'Mesmo Relatório')
doc.add_page_break()
add_analysis_to_doc(doc, stats_cross, 'Fatura x Caixa')

doc.add_page_break()
doc.add_heading('Detalhamento Completo', level=1)
doc.add_paragraph(f'Os {total} pares de duplicatas estão listados no arquivo Excel anexo.')
doc.add_paragraph('')

doc.add_heading('Top 50 Duplicatas por Valor', level=2)
sorted_pairs = sorted(pairs, key=lambda p: p['original']['value'], reverse=True)[:50]
table = doc.add_table(rows=len(sorted_pairs) + 1, cols=7, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['ID Orig.', 'Título', 'Valor', 'Data', 'Usuário', 'ID Dup.', 'Mesmo Rel.']
for j, h in enumerate(headers):
    cell = table.rows[0].cells[j]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(8)
for i, p in enumerate(sorted_pairs, 1):
    o = p['original']
    d = p['duplicate']
    row = table.rows[i]
    row.cells[0].text = str(o['expense_id'])
    row.cells[1].text = o['title'][:30]
    row.cells[2].text = f'R$ {o["value"]:,.2f}'
    row.cells[3].text = o['date']
    row.cells[4].text = o['user_name'][:20]
    row.cells[5].text = str(d['expense_id'])
    row.cells[6].text = 'Sim' if d['same_report'] else 'Não'
    for cell in row.cells:
        for pp in cell.paragraphs:
            for r in pp.runs:
                r.font.size = Pt(8)

docx_path = 'duplicatas_2026_report.docx'
doc.save(docx_path)
print(f'DOCX saved: {docx_path}')
print(f'Total pairs: {total}')
print(f'Same report: {len(same_report)} (R$ {stats_same["total_value"]:,.2f})')
print(f'Diff report: {len(diff_report)} (R$ {stats_diff["total_value"]:,.2f})')
print(f'Fatura x Caixa: {len(cross_fatura_caixa)} (R$ {stats_cross["total_value"]:,.2f})')
